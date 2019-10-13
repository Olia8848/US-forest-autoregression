# -*- coding: utf-8 -*-
"""
Created on Wed Jul 17 15:22:20 2019

@author: olga
"""
# pip install python-docx

import numpy
import math
import scipy
import pandas   
import matplotlib.pyplot as plt
import random
import docx

from numpy import random
from numpy import mean
from matplotlib import pyplot
from statsmodels.graphics.gofplots import qqplot
from scipy import stats

from docx import Document
from docx.shared import Inches   
from docx.shared import Pt   


def Ev(a, n):
    if (a == -1) or (a == 1):
        return math.sqrt(n)
    else:
        b = ( 1-a**(2*n) ) / ( 1-a**2 )  
        return math.sqrt(b)

def Od(a, n): 
    if (a == 1):
        return n
    else:
        return (1-a**(n))/(1-a)   
     
# This is our regression with given value of a 
def regression(a, logvalue, lognext, interval, NCHANGES):
    V = numpy.array([])
    A = numpy.array([])
    for j in range(NCHANGES): 
         d = interval[j]
         diff = lognext[j]-(a**d)*logvalue[j]
         V = numpy.append(V, diff/Ev(a, d))  
         A = numpy.append(A, Od(a, d)/Ev(a, d))
    r = numpy.dot(V, A)/numpy.dot(A, A)
    sigma = math.sqrt((1/(NCHANGES-1))*numpy.dot(V-r*A, V-r*A))
    return V, A, r, sigma

# QQ plot of residuals for this regression
# compare to standard normal distribution, 'r' - regression line    
def qq(a, logvalue, lognext, interval, NCHANGES, im, eid):
    V, A, r, sigma = regression(a, logvalue, lognext, interval, NCHANGES)
    centralized = numpy.array([V[k] - r*A[k] for k in range(NCHANGES)])
    s1 = stats.shapiro(centralized)[0]
    s2 = stats.shapiro(centralized)[1]
    fig = qqplot(centralized, line = 'r') 
    #fig = plt.figure()
    fig.savefig(im + 'eco_'+ Ecoregions[eid] + '_Raw_Data_residuals.png')
    pyplot.show()
    return s1, s2


def inverseGamma(alpha, beta):
    return (1/numpy.random.gamma(alpha, 1/beta))



def Histograms(meansB, varsB, t, im, eid):
    msB = numpy.array([])
    vsB = numpy.array([])

    for j in range(NSIMS):
        msB = numpy.append(msB, meansB[t][j])      
        vsB = numpy.append(vsB, varsB[t][j])
    time = Years[t]

    im = 'C:/Users/olga/Desktop/US_forest/pictures/'

    fig = plt.figure()
    plt.hist(msB, bins = 50) 
    plt.title("simulated means \n of ln(biomass) in " + str(time))
#    plt.title("simulated means \n of ln(BA) in " + str(time))
    plt.show()
    fig.savefig(im + 'eco_'+ Ecoregions[eid] + 'simulated means of ln(biomass) in' + str(time) + '.png')
    results.add_picture(im + 'eco_'+ Ecoregions[eid] + 'simulated means of ln(biomass) in' + str(time) + '.png', width = Inches(3.0))
 
    fig = plt.figure()
    plt.hist(vsB, bins = 50) 
    plt.title("simulated variances \n of ln(biomass) in " + str(time))
#    plt.title("simulated variances \n of ln(BA) in " + str(time))
    plt.show()
    fig.savefig(im + 'eco_'+ Ecoregions[eid] + 'simulated variances of ln(biomass) in' + str(time) + '.png')
    results.add_picture(im + 'eco_'+ Ecoregions[eid] + 'simulated variances of ln(biomass) in' + str(time) + '.png', width = Inches(3.0))
  





path = 'C:/Users/olga/Desktop/US_forest/'
# path = '/Users/Olga Rumyantseva/Desktop/Python biomass/'

im = 'C:/Users/olga/Desktop/US_forest/pictures/'


df = pandas.read_csv(path + 'biodataUS_sorted.csv')


df.columns

df1 = df.iloc[:, [2, 3, 18, 19, 24, 25, 22, 23]]
df1.columns


data0 = df1.values  
NOBSERV0 = numpy.size(data0, 0) 
print(NOBSERV0) #  number of observations totally 409 868



patch = data0[:,0]   # Plot ID 
year = numpy.array([int(data0[k,1]) for k in range(NOBSERV0)]) # Year 
biomass = data0[:,2] # Biomass
barea = data0[:,3]   # Basal Area
ecoreg = numpy.array([str(data0[k,4]) for k in range(NOBSERV0)])  # Eco region
state = numpy.array([str(data0[k,5]) for k in range(NOBSERV0)])   # US State


patch[0]   # Plot ID
year[0]    # Year
biomass[0] # Biomass
barea[0]   # Basal Area
ecoreg[0]  # Eco region
state[0]   # US State


Ecoregions  = numpy.unique(ecoreg)
NEcoregions = numpy.size(Ecoregions) # 36 ecoregions

States = numpy.unique(state)
NStates = numpy.size(States)  # 48 states

Years  = numpy.unique(year)
NYears = numpy.size(Years) # 40 years: 1968 - 2013
# 1968, 1970, 1972, 1974, 1977, 1978, 1980, 1981, 1982, 1983, 1984,
#       1985, 1986, 1987, 1988, 1989, 1990, 1991, 1992, 1993, 1994, 1995,
#       1996, 1997, 1998, 1999, 2000, 2001, 2002, 2003, 2004, 2005, 2006,
#      2007, 2008, 2009, 2010, 2011, 2012, 2013


###########################################################################
#########   Restricted by ecoregion: ######################################
###########################################################################

results = docx.Document() 

eid = 0

df2 = df1[ecoreg == Ecoregions[eid]]

results.add_heading('Ecoregion ' + Ecoregions[eid])

data = df2.values  

NOBSERV = numpy.size(data, 0) 

print(NOBSERV)  # observations in your ecoregion

results.add_paragraph('Number of observations in the ecoregion is ' + str(NOBSERV))


patch = data[:,0]   # Plot ID 
year = numpy.array([int(data[k,1]) for k in range(NOBSERV)]) # Year 
biomass = data[:,2] # Biomass
barea = data[:,3]   # Basal Area
ecoreg = numpy.array([str(data[k,4]) for k in range(NOBSERV)])  # Eco region
state = numpy.array([str(data[k,5]) for k in range(NOBSERV)])   # US State
# value = biomass # this is what we consider now: biomass 
value = biomass # this is what we consider now: biomass 




logvalue = numpy.array([]) # here will be logarithms of value for steps
lognext = numpy.array([])  # increments of logarithms 
interval = numpy.array([]) # increments of years

#next loop collects records from the same patch: 
#initial , change in logs, time interval 
#For example, if the same patch is observed in 1980 and 1987, 
#collect logvalue in 1980, logchange and time interval

for observ in range(NOBSERV-1):
    if patch[observ] == patch[observ+1]:
        logvalue = numpy.append(logvalue, math.log(value[observ]))
        lognext = numpy.append(lognext, math.log(value[observ+1]))
        interval = numpy.append(interval, year[observ+1] - year[observ])

       
#Number of pair-observations
NCHANGES = numpy.size(logvalue)  
print('Number of pair-observations = ', NCHANGES)


results.add_paragraph('Number of pair-observations = ' + str(NCHANGES))


Years = numpy.unique(year)
NYears = numpy.size(Years)


results.add_paragraph('Number of years = ' + str(NYears))
results.add_paragraph('Years : ' + str(Years))



States = numpy.unique(state)
NStates = numpy.size(States) 


results.add_paragraph('Number of years observed = ' + str(NStates))
results.add_paragraph('Years observed: ' + str(States))




#Average biomass or shadow in a given year
YearAvBiomass = numpy.array([])
for j in range(NYears):
    logbio = math.log(numpy.mean(value[year == Years[j]]))
    YearAvBiomass = numpy.append(YearAvBiomass, logbio)
    
#Increments of average biomass (or shadow)
DeltasYearAvBiomass = numpy.array([])  
for i in range(numpy.size(YearAvBiomass)-1):
    delt = YearAvBiomass[i+1]-YearAvBiomass[i]
    DeltasYearAvBiomass = numpy.append(DeltasYearAvBiomass, delt)  
    

########################################################################
########################################################################
########################################################################
########################################################################
########################################################################
    
    
value = barea # this is what we consider now
logvalue1 = numpy.array([]) # here will be logarithms of value for steps
lognext1 = numpy.array([])  # increments of logarithms 
interval1 = numpy.array([]) # increments of years

for observ in range(NOBSERV-1):
    if patch[observ] == patch[observ+1]:
        logvalue1 = numpy.append(logvalue1, math.log(value[observ]))
        lognext1 = numpy.append(lognext1, math.log(value[observ+1]))
        interval1 = numpy.append(interval1, year[observ+1] - year[observ])

YearAvBiomass1 = numpy.array([])
for j in range(NYears):
    logbio = math.log(numpy.mean(value[year == Years[j]]))
    YearAvBiomass1 = numpy.append(YearAvBiomass1, logbio)
          


c1 = numpy.correlate(logvalue, logvalue1)
c2 = numpy.correlate(YearAvBiomass, YearAvBiomass1)


results.add_paragraph('Correlation between ln(biomass) and ln(basal area): ' + str(c1))

results.add_paragraph('Correlation between ln(year average biomass) and ln(year average basal area): ' + str(c2))



print('mean change of avg year biomass = ', numpy.mean(DeltasYearAvBiomass))
#Mean of increments of average yearly biomass
print('stdev of this change = ', numpy.std(DeltasYearAvBiomass))
#Standard deviation of these increments


results.add_paragraph('Mean change of year avg. biomass = ' + str(numpy.mean(DeltasYearAvBiomass)))

results.add_paragraph('Std. of the change of year avg. biomass = ' + str(numpy.std(DeltasYearAvBiomass)))

results.save(path + 'eco_' + Ecoregions[eid] + '_results.docx')


########################################################################
########################################################################
########################################################################
########################################################################
########################################################################
    

value = biomass # this is what we consider now

Sigmas = numpy.array([])    
Rs = numpy.array([])

x = numpy.arange(-2, 2, 0.01)
for a in x:
    V, A, r, sigma = regression(a, logvalue, lognext, interval, NCHANGES)
    Sigmas = numpy.append(Sigmas, sigma)
    Rs = numpy.append(Rs, r)

#Plot of parameter a vs standard error
    
fig = plt.figure()
pyplot.plot(x, Sigmas)
pyplot.show()
fig.savefig(im + 'eco_'+ Ecoregions[eid] + '_Raw_Data_a_vs_stand_error.png')

results.add_picture(im + 'eco_'+ Ecoregions[eid] + '_Raw_Data_a_vs_stand_error.png', width = Inches(3.0))


r = Rs[numpy.argmin(Sigmas)]#Corresponding r

c1 = numpy.min(Sigmas)
c2 = x[numpy.argmin(Sigmas)]


print('Non-centered data for all years')
print('min std = ', c1)
print('value of a = ', c2)
print('value of m = ', r)


results.add_paragraph('Non-centered data for all years: ')


results.add_paragraph('Min. Std. = ' + str(c1))
results.add_paragraph('Value of a = '+ str(c2))
results.add_paragraph('Value of m = '+ str(r))




print('qq plot for residuals for random walk case')


qq(1, logvalue, lognext, interval, NCHANGES, im, eid)


#residuals = [(lognext[k] - logvalue[k])/math.sqrt(interval[k]) - m * math.sqrt(interval[k]) for k in range(NCHANGES)]

results.add_picture(im + 'eco_'+ Ecoregions[eid] + '_Raw_Data_residuals.png', width = Inches(2.0))


########################################################################
########################################################################
########################################################################
########################################################################
########################################################################


mu = numpy.mean(DeltasYearAvBiomass)   # mean   numpy.mean(DeltasYearAvBiomass)
sigma = numpy.std(DeltasYearAvBiomass) # standard deviation   numpy.std(DeltasYearAvBiomass)

from math import e

# Years = numpy.unique(year)
# m0 = numpy.mean(value[year == Years[0]]  # Years[0] == 1970

#  means and variances based on existing data:
ExpectationsB = numpy.array([]) 
VariancesB = numpy.array([])
ExpectationsBA = numpy.array([]) 
VariancesBA = numpy.array([])

results.add_paragraph('   ')


#  means and variances based on existing data:
for t in range(NYears): 
       meanB = numpy.mean(biomass[year == Years[t]])
       meanB = round(meanB, 2)
       ExpectationsB = numpy.append(ExpectationsB, meanB)
       
       vB = numpy.var(biomass[year == Years[t]])
       vB = round(math.sqrt(vB), 2)
       VariancesB = numpy.append(VariancesB, vB)
       
       meanBA = numpy.mean(barea[year == Years[t]])
       meanBA = round(meanBA, 2)
       ExpectationsBA = numpy.append(ExpectationsBA, meanBA)
       
       vBA = numpy.var(barea[year == Years[t]])
       vBA = round(math.sqrt(vBA), 2)
       VariancesBA = numpy.append(VariancesBA, vBA)
       
       nobserv = numpy.size(biomass[year == Years[t]])
       
       print(Years[t],', nobserv.:', nobserv, ', meanB:', meanB, ', vB:', vB,', meanBA:',  meanBA, ', vBA:', vBA, '\\')
       results.add_paragraph(str(Years[t]) + ',  nobserv.:  ' + str(nobserv) + ',  meanB:  ' + str(meanB) + ',  varB:  ' + str(vB) + ',  meanBA:  ' +  str(meanBA) + ',  varBA:  ' + str(vBA)) 


#  simulated means and variances:
ExpB = numpy.array([]) 
VarB = numpy.array([])
ExpBA = numpy.array([]) 
VarBA = numpy.array([])



for t in range(2021 - Years[NYears-1]): 
       m0 = numpy.mean(biomass[year == Years[NYears-1]])
       meanB = round(m0*(e**(t*mu + t*(sigma**2)*(1/2))), 1)
       meanB = round(meanB, 2)
       ExpB = numpy.append(ExpB, meanB)
       
       vB = (m0**2)*(e**(2*t*mu + t*(sigma**2)*(1/2)))*(e**(t*(sigma**2)) - 1)
       vB = round(math.sqrt(vB), 2)
       VarB = numpy.append(VarB, vB)
       
       m0 = numpy.mean(barea[year == Years[0]])
       meanBA = round(m0*(e**(t*mu + t*(sigma**2)*(1/2))), 2)
       meanBA = round(meanBA, 2)
       
       ExpBA = numpy.append(ExpBA, meanBA)
       vBA = (m0**2)*(e**(2*t*mu + t*(sigma**2)*(1/2)))*(e**(t*(sigma**2)) - 1)
       vBA = round(math.sqrt(vBA), 1)
       VarBA = numpy.append(VarBA, vBA)
       
       print(Years[NYears-1]+t,', meanB:',  meanB, ', vB:', vB,', meanBA:',  meanBA, ', vBA:', vBA, '\\')
       results.add_paragraph(str(Years[NYears-1]+t) + '  , meanB:'  +  str(meanB) + ',  vB: ' + str(vB) +',  meanBA: ' +  str(meanBA) + ' , vBA: ' + str(vBA)) 



########################################################################
########################################################################
########################################################################
########################################################################
########################################################################

# Annual means of biomass:
AnnualMeanBiomass = numpy.array([]) 

for i in range(NYears):
    logbio = math.log(numpy.mean(value[year == Years[i]]))
    AnnualMeanBiomass = numpy.append(AnnualMeanBiomass, logbio)
    

# Annual means of shadows:
AnnualMeanBA = numpy.array([]) 

for i in range(NYears):
    logBA = math.log(numpy.mean(barea[year == Years[i]]))
    AnnualMeanBA = numpy.append(AnnualMeanBA, logBA)
    

##################################################################
#####  Plots of annual means of biomass and shadow: ###############
##################################################################

fig1 = plt.figure()
pyplot.plot(Years, AnnualMeanBiomass, 'bo', Years, AnnualMeanBiomass, 'k')
pyplot.show()
fig1.savefig(im + 'eco_' + Ecoregions[eid] + '_AnnualMeanBiomass.png')

results.add_picture(im + 'eco_'+ Ecoregions[eid] + '_AnnualMeanBiomass.png', width = Inches(3.0))


fig2 = plt.figure()
pyplot.plot(Years, AnnualMeanBA, 'bo', Years, AnnualMeanBA, 'k')
pyplot.show()
fig2.savefig(im + 'eco_'+ Ecoregions[eid] + '_AnnualMeanBA.png')

results.add_picture(im + 'eco_'+ Ecoregions[eid] + '_AnnualMeanBA.png', width = Inches(3.0))

  
########################  Table (38*4)  ###########################

# Year|| 
# of observ ||
# Mean of biom. this year || 
# Mean of shadow this year  ||

####################################################################

NObserv = [0] * NYears 

for i in range(NYears):
    NObserv[i] = numpy.size(value[year == Years[i]])

dataT3 = {'Year':Years,'Number of Observations':NObserv, 
'mean of biomass':AnnualMeanBiomass,'mean of basal area':AnnualMeanBA}
df3 = pandas.DataFrame(dataT3)
print(df3)

df3.to_csv(path + 'eco_'+ Ecoregions[eid] + '_means_biomass_BA.csv', index=False)


##################################################################
#####  Distribution of patches observed in 2007: ###############
##################################################################

# the last year for which we have observations is Years[NYears-1]

PatchesLastObsYear = patch[year == Years[NYears-1]]
BiomassLastObsYear = biomass[year == Years[NYears-1]]
BALastObsYear = barea[year == Years[NYears-1]]


# pyplot.hist(biomass[year == 2012] , bins = 100)
# pyplot.show()

# pyplot.hist(BALastObsYear, bins = 100)
# pyplot.show()

##################################################################
#####  Simulate Random Walk 12 years ahead from 2007
# for some patch which was observed in 2007: ###############
##################################################################
Iteartions = 1000  # the number of simulations 


# for ex., 1000 simulations of 12-year forward predictions:
# select randomly patch which was observed in 2007
# (Y0,Y2, ... , Y11) ... (Y0,Y2, ... , Y11) - 1000 arrays like this
# Y0 - in 2008 (1st prediction), ... , Y11 - in 2019

n = 2022 - Years[NYears-1]
 
Simulations = [[0] * n for i in range(Iteartions)] 

PatchesLastObsYear = patch[year == Years[NYears-1]]

for i in range(Iteartions):
    dat1 = data[(year == Years[NYears-1]) & (patch == random.choice(PatchesLastObsYear))]
    value = dat1[0][2] # biomass
    Simulations[i][0] = value + numpy.random.normal(mu, sigma) # 2008 prediction
#   Switch between these two:
#   value = dat1[0][2] # biomass on the chosen patch in 2007
#   value = dat1[0][3] # bas. area on the chosen patch in 2007

Simulations[100][0]

for j in range(n-1):
    print('year =', Years[NYears-1]+j+2) # doing simulation for this year
    for i in range(Iteartions):
        Simulations[i][j+1] = Simulations[i][j] + numpy.random.normal(mu, sigma) 
    

    
Predictions = numpy.array([]) 

for j in range(n):
    sim = numpy.array([])
    for i in range(Iteartions):
        sim = numpy.append(sim, Simulations[i][j])
    Predictions = numpy.append(Predictions, math.log(numpy.mean(sim)))


YearsPred = numpy.array([])

for j in range(n):
    YearsPred = numpy.append(YearsPred, Years[NYears-1] + j+1)


# fig, ax = pyplot.subplots()
# pyplot.plot(YearsPred, Predictions, 'bo', YearsPred, Predictions, 'k', color='tab:green')
# ax.ticklabel_format(useOffset=False)
# pyplot.plot(Years, AnnualMeanBiomass, 'bo', Years, AnnualMeanBiomass, 'k')
# plt.show()



        

########################################################################
########################################################################
########################################################################
########################################################################
########################################################################

Patches  = numpy.unique(patch)
NPatches = numpy.size(Patches)

value = biomass

LogValuePatchYear = [[0] * NPatches for i in range(NYears)] 
# [[p1,..., pm]  [p1,..., pm] .... [p1,..., pm]]
#     year_1       year_2             year_n


#   LogValuePatchYear[i][p] = log(biomass at plot p in year i): 
for k in range(NOBSERV):
    indY = numpy.where(Years == year[k])
    i = indY[0][0]     # return i such that year[k] == Years[i]
    indP = numpy.where(Patches == patch[k])
    j = indP[0][0]  # return j such that patch[k] == Patches[j]
    LogValuePatchYear[i][j] = math.log(value[k])
#   i - year, j - patch

#  LogValuePatchYear[year][patch]
 
############################################################################
############################################################################
############################################################################  

NSIMS = 1000 # Number of simulations of Bayesian estimates

# means simulated by Bayes:
meansB =  [[0] * NSIMS for i in range(NYears)] 
# [[m_1,..., m_1000]  .... [m_1,..., m_1000]]
#     year_1                  year_38
#  meansB[year][simulation]


# vars simulated by Bayes:
varsB = [[0] * NSIMS for i in range(NYears)]
#  varsB[year][simulation]


for t in range(NYears):    
    LogValuesYear = numpy.array([])  # biomasses vector x_{1}(t), ... , x_{#p}(t)
    for p in range(NPatches):    
        LogValuesYear = numpy.append(LogValuesYear, LogValuePatchYear[t][p])
    LogValuesYear = LogValuesYear[numpy.nonzero(LogValuesYear)] #  biomass data doesn't  contain 0 records, so we can remove zeroes:
#    print('year = ', t)
    empMean = numpy.mean(LogValuesYear) # mean of biomasses in year t
    empVar = numpy.var(LogValuesYear) # variance of biomasses in year t
#    print('mean = ', round(empMean, 2))
#    print('var = ', round(empVar, 2))
    n = len(LogValuesYear) # number of patches observed in year t
#    print('number = ', n)
    if n == 1:
       continue
    print(Years[t], ' & number of patches observed in this year: ', n, ' & empir. mean: ', round(empMean, 2), ' & empir. variance:', round(empVar, 2))
    results.add_paragraph(str(Years[t]) + '  & number of patches observed in this year:  ' + str(n) + '  & empir. mean:  ' + str(round(empMean, 2)) + '  & empir. variance: ' + str(round(empVar, 2))) 
    for j in range(NSIMS):
        varsB[t][j] = inverseGamma((n-1)/2, (n*empVar)/2) 
        meansB[t][j] = random.normal(empMean, varsB[t][j]/n)



Histograms(meansB, varsB, 0, im, eid)
        
         
# g^hat - estimate:
summ = 0
for i in range(NSIMS):
    summ = summ + meansB[NYears-1][i]-meansB[0][i]  

empG = summ/(NSIMS*NYears)  # g^{hat} - average growth rate
print('global mean = ', empG)
results.add_paragraph('Global Mean = ' + str(empG)) 

# meansB[t][i]-meansB[t-1][i]    

# sigma^hat - estimate:
summ = 0
for i in range(NSIMS):
    for t in range(NYears):
        summ = summ + (meansB[t][i]-meansB[t-1][i] - empG)**2

empGlobalV = summ/(NSIMS*NYears)   
print('global stdev = ', empGlobalV) 
results.add_paragraph('Global Stdev. = ' + str(empGlobalV)) 
 

# G, Sigma - estimates, updated by Bayes:
Sigma = inverseGamma((NSIMS*NYears-1)/2, (NSIMS*NYears*empGlobalV)/2)
G = random.normal(empG, Sigma/(NSIMS*NYears))
results.add_paragraph('Bayes Sigma = ' + str(Sigma)) 
results.add_paragraph('Bayes G = ' + str(G)) 
    
    
results.save(path + 'eco_' + Ecoregions[eid] + '_results.docx')
